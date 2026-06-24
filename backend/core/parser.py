import io
import re
from backend.utils.log import logger

ALLOWED_MIME_TYPES = {
    "application/pdf": b"%PDF",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": b"PK\x03\x04",
}


def validate_file_bytes(file_bytes: bytes, filename: str) -> str:
    for mime, magic in ALLOWED_MIME_TYPES.items():
        if file_bytes[:len(magic)] == magic:
            return mime
    raise ValueError(f"Unsupported file type for '{filename}'. Only PDF and DOCX accepted.")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text(layout=True)
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def clean_extracted_text(text: str) -> str:
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    lines = text.split('\n')
    line_counts = {}
    for line in lines:
        stripped = line.strip()
        if len(stripped) > 10:
            line_counts[stripped] = line_counts.get(stripped, 0) + 1
    repeated = {line for line, count in line_counts.items() if count >= 3}
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if stripped in repeated:
            continue
        if re.match(r'^\d+\s*/\s*\d+$', stripped):
            continue
        if re.match(r'^Seite\s+\d+', stripped, re.IGNORECASE):
            continue
        if re.match(r'^Page\s+\d+', stripped, re.IGNORECASE):
            continue
        if re.match(r'^\d{1,2}\.\d{1,2}\.\d{2,4}$', stripped):
            continue
        cleaned.append(line)
    return '\n'.join(cleaned)


def extract_keywords_from_text(text: str) -> list[str]:
    keywords = set()
    tech_pattern = re.compile(
        r'\b(?:Python|Java|JavaScript|TypeScript|SQL|NoSQL|React|FastAPI|Django|Flask|'
        r'Docker|Kubernetes|AWS|Azure|GCP|Linux|Git|JIRA|SAP|Oracle|Excel|'
        r'PACS|HIS|KIS|RIS|Medicom|iMedOne|Orbis|EHR|EMR|HL7|FHIR|'
        r'ICD-10|ICD-11|DRG|CT|MRI|PET|ECG|EEG|ICU|ER|OR)\b',
        re.IGNORECASE
    )
    keywords.update(m.group() for m in tech_pattern.finditer(text))
    cert_pattern = re.compile(
        r'\b(?:Dr\.?|Prof\.?|M\.D\.?|Ph\.D\.?|MBA|MSc|BSc|'
        r'Approbation|Berufserlaubnis|Weiterbildung|Fachsprachprüfung|'
        r'[A-Z]{2,6}(?:\s[A-Z]{2,6})?)\b'
    )
    keywords.update(m.group() for m in cert_pattern.finditer(text))
    lang_pattern = re.compile(r'\b(?:B1|B2|C1|C2|A1|A2|TestDaF|DSH|IELTS|TOEFL|Goethe)\b')
    keywords.update(m.group() for m in lang_pattern.finditer(text))
    title_pattern = re.compile(r'\b[A-ZÜÖÄ][a-züöäß]+(?:\s[A-ZÜÖÄ][a-züöäß]+){1,3}\b')
    skip = {"January", "February", "March", "April", "August", "September",
            "October", "November", "December", "Januar", "Februar", "März", "Oktober", "Dezember"}
    for match in title_pattern.finditer(text):
        phrase = match.group()
        if len(phrase) > 5 and phrase not in skip:
            keywords.add(phrase)
    cleaned = [
        k.strip() for k in keywords
        if len(k.strip()) > 2 and not k.strip().isdigit()
        and k.strip() not in {"The", "And", "For", "With", "Und", "Mit", "Der", "Die", "Das"}
    ]
    logger.debug(f"Parser extracted {len(cleaned)} raw keywords")
    return sorted(set(cleaned))


def parse_file(file_bytes: bytes, filename: str) -> tuple[str, list[str], str]:
    mime_type = validate_file_bytes(file_bytes, filename)
    if mime_type == "application/pdf":
        text = extract_text_from_pdf(file_bytes)
    else:
        text = extract_text_from_docx(file_bytes)
    if not text.strip():
        raise ValueError("Could not extract text from file. May be image-based or corrupted.")
    text = clean_extracted_text(text)
    keywords = extract_keywords_from_text(text)
    logger.info(f"Parsed '{filename}': {len(text)} chars, {len(keywords)} raw keywords")
    return text, keywords, mime_type
